"""Report generation tools for converting Markdown to various formats."""

import logging
import math
import os
import re
from datetime import datetime
from pathlib import Path

from langchain.tools import ToolRuntime
from langchain_core.messages import ToolMessage
from langgraph.types import Command

from .config import get_enabled_tools

logger = logging.getLogger(__name__)


def get_report_tools() -> list:
    """Get report generation tools.
    
    Returns:
        List of report generation tools.
        
    Available tools:
        - markdown_to_html_report: Convert Markdown to interactive HTML with Chart.js
        - markdown_to_docx: Convert Markdown to Word document
    """
    enabled = get_enabled_tools()
    tools = []
    tool_names = []
    
    # ============================================================
    # Tool 1: Markdown to HTML Report
    # ============================================================
    if enabled.get("markdown_to_html_report", True):
        try:
            import markdown
            from markdown.extensions.tables import TableExtension
            from markdown.extensions.fenced_code import FencedCodeExtension
            
            def markdown_to_html_report(
                markdown_content: str,
                title: str = "Report",
                runtime: ToolRuntime = None,
            ) -> Command | dict:
                """Convert Markdown to a beautiful, professional HTML report.
                
                IMPORTANT GUIDELINES FOR INPUT MARKDOWN:
                =====================================
                
                1. **USE TABLES EXTENSIVELY** - Present all data, comparisons, and metrics 
                   using well-structured Markdown tables. Tables will be rendered beautifully
                   with professional styling (zebra stripes, hover effects, header colors).
                   
                2. **TABLE STRUCTURE** - Each table should have:
                   - Clear, concise column headers
                   - Properly aligned data
                   - Consistent formatting (numbers, percentages, dates)
                   - **CRITICAL**: Blank lines BEFORE and AFTER the table
                   
                   ✅ CORRECT Example:
                   
                   **Competitive Analysis**
                   
                   | Metric | Value | Change |
                   |--------|-------|--------|
                   | Traffic | 50,000 | +15% |
                   | Keywords | 1,200 | +8% |
                   
                   Next paragraph starts here...
                   
                   ❌ WRONG Example (no blank lines):
                   **Competitive Analysis**
                   | Metric | Value | Change |
                   |--------|-------|--------|
                   | Traffic | 50,000 | +15% |
                   Next paragraph starts here...
                   
                   Without blank lines, the table will NOT render correctly!
                
                3. **VISUAL HIERARCHY** - Use proper heading levels (h1 > h2 > h3)
                   to create clear document structure.
                
                4. **CHARTS** - Tables are automatically converted to charts:
                   - Time-series tables (dates/months) → Line charts
                   - Comparison tables → Bar charts  
                   - Percentage tables → Pie charts
                
                5. **FORMATTING** - Use bold (**text**) for emphasis, 
                   code blocks for technical content.
                
                Args:
                    markdown_content: Well-structured Markdown with tables for data presentation
                    title: Report title (appears in header and browser tab)
                    runtime: Tool runtime for state management (automatically provided)
                
                Returns:
                    Command object to update artifacts, or dict with error message
                """
                try:
                    # Convert Markdown to HTML
                    md = markdown.Markdown(
                        extensions=[
                            TableExtension(),
                            FencedCodeExtension(),
                            'extra',
                            'codehilite',
                            'toc'
                        ]
                    )
                    html_content = md.convert(markdown_content)
                    
                    # Extract tables for chart generation
                    tables = _extract_tables(markdown_content)
                    chart_configs = _generate_chart_configs(tables)
                    
                    # Generate complete HTML with styling and Chart.js
                    full_html = _generate_html_template(
                        title=title,
                        content=html_content,
                        chart_configs=chart_configs
                    )
                    
                    # Save to file - both in workspace root (for artifacts) and reports/ (for archive)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"report_{timestamp}.html"
                    
                    # Get workspace root directory (parent of deepagents/)
                    workspace_root = Path(__file__).parent.parent.parent
                    
                    # Save to workspace root directory for artifacts display
                    root_filepath = workspace_root / filename
                    with open(root_filepath, 'w', encoding='utf-8') as f:
                        f.write(full_html)
                    
                    # Also save to reports/ directory for archiving
                    output_dir = Path("reports")
                    output_dir.mkdir(exist_ok=True)
                    archive_filepath = output_dir / filename
                    with open(archive_filepath, 'w', encoding='utf-8') as f:
                        f.write(full_html)
                    
                    filepath = root_filepath  # Use root path for return value

                    # Add to artifacts if runtime is provided
                    if runtime is not None:
                        # Update LangGraph state with the new file
                        files = runtime.state.get("files", {})
                        files_update = {filename: {"content": full_html}}
                        
                        # Prepare structured result with HTML content
                        import json
                        result_data = {
                            "success": True,
                            "title": title,
                            "filename": filename,
                            "file_path": str(filepath.absolute()),
                            "html_content": full_html,
                            "charts_generated": len(chart_configs),
                            "file_size": len(full_html),
                            "message": f"✅ HTML report generated: {filename}\n\n"
                                      f"📊 Charts: {len(chart_configs)} interactive visualizations\n"
                                      f"📄 Size: {len(full_html):,} bytes\n"
                                      f"🎨 Professional design with sidebar navigation\n\n"
                                      f"The report is now available in the Artifacts panel."
                        }
                        
                        return Command(
                            update={
                                "files": files_update,
                                "messages": [
                                    ToolMessage(
                                        content=json.dumps(result_data, ensure_ascii=False),
                                        tool_call_id=runtime.tool_call_id,
                                    )
                                ],
                            }
                        )
                    
                    # Fallback for when runtime is not available
                    return {
                        "success": True,
                        "title": title,
                        "filename": filename,
                        "file_path": str(filepath.absolute()),
                        "html_content": full_html,
                        "charts_generated": len(chart_configs),
                        "file_size": len(full_html),
                        "message": f"HTML report generated: {filename}",
                    }
                except Exception as e:
                    logger.error(f"[REPORTS] Error generating HTML report: {e}")
                    return {
                        "success": False,
                        "error": str(e),
                        "message": "Failed to generate HTML report"
                    }
            
            tools.append(markdown_to_html_report)
            tool_names.append("markdown_to_html_report")
            
        except ImportError as e:
            logger.warning(f"[REPORTS] Required packages not installed for HTML reports: {e}")
            logger.warning("[REPORTS] Run: pip install markdown")
    
    # ============================================================
    # Tool 2: Markdown to DOCX
    # ============================================================
    if enabled.get("markdown_to_docx", True):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.style import WD_STYLE_TYPE
            
            def markdown_to_docx(
                markdown_content: str,
                filename: str = "report.docx",
                runtime: ToolRuntime = None,
            ) -> Command | dict:
                """Convert Markdown to Word document (.docx).
                
                Args:
                    markdown_content: The Markdown content to convert
                    filename: Output filename (default "report.docx")
                    runtime: Tool runtime for state management (automatically provided)
                
                Returns:
                    Command object to update artifacts, or dict with error message
                """
                try:
                    document = Document()
                    
                    # Set document styles
                    _setup_docx_styles(document)
                    
                    # Parse and add content
                    lines = markdown_content.split('\n')
                    i = 0
                    while i < len(lines):
                        line = lines[i]
                        
                        # Headers
                        if line.startswith('# '):
                            document.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            document.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            document.add_heading(line[4:], level=3)
                        elif line.startswith('#### '):
                            document.add_heading(line[5:], level=4)
                        
                        # Tables
                        elif line.startswith('|') and '|' in line:
                            table_lines, i = _extract_table_lines(lines, i)
                            if table_lines:
                                _add_table_to_docx(document, table_lines)
                            i -= 1  # Adjust for increment
                        
                        # Code blocks
                        elif line.startswith('```'):
                            code_lines, i = _extract_code_block(lines, i)
                            if code_lines:
                                p = document.add_paragraph('\n'.join(code_lines))
                                p.style = 'Code'
                            i -= 1
                        
                        # Blockquotes
                        elif line.startswith('> '):
                            document.add_paragraph(line[2:], style='Quote')
                        
                        # Lists
                        elif line.strip().startswith('- ') or line.strip().startswith('* '):
                            document.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif re.match(r'^\d+\.\s', line.strip()):
                            text = re.sub(r'^\d+\.\s', '', line.strip())
                            document.add_paragraph(text, style='List Number')
                        
                        # Regular paragraphs
                        elif line.strip():
                            # Remove Markdown formatting
                            clean_text = _clean_markdown(line)
                            document.add_paragraph(clean_text)
                        
                        i += 1
                    
                    # Save document
                    output_dir = Path("reports")
                    output_dir.mkdir(exist_ok=True)
                    
                    if not filename.endswith('.docx'):
                        filename += '.docx'
                    
                    filepath = output_dir / filename
                    document.save(str(filepath))
                    
                    file_size = filepath.stat().st_size
                    
                    # DOCX files can't be displayed in artifacts (binary format)
                    # But we notify the user via ToolMessage
                    if runtime is not None:
                        return Command(
                            update={
                                "messages": [
                                    ToolMessage(
                                        content=f"✅ Word document generated: {filename}\n\n"
                                                f"📁 Location: {filepath.absolute()}\n"
                                                f"📄 Size: {file_size:,} bytes\n\n"
                                                f"The document has been saved to the reports/ directory.",
                                        tool_call_id=runtime.tool_call_id,
                                    )
                                ],
                            }
                        )
                    
                    # Fallback for when runtime is not available
                    return {
                        "success": True,
                        "filename": filename,
                        "file_path": str(filepath),
                        "file_url": f"file://{filepath.absolute()}",
                        "file_size": file_size,
                        "message": f"Word document generated: {filename}",
                    }
                except Exception as e:
                    logger.error(f"[REPORTS] Error generating DOCX: {e}")
                    return {
                        "success": False,
                        "error": str(e),
                        "message": "Failed to generate Word document"
                    }
            
            tools.append(markdown_to_docx)
            tool_names.append("markdown_to_docx")
            
        except ImportError as e:
            logger.warning(f"[REPORTS] Required packages not installed for DOCX generation: {e}")
            logger.warning("[REPORTS] Run: pip install python-docx")
    
    if tools:
        logger.info(f"[REPORTS] Tools enabled: {', '.join(tool_names)}")
    else:
        logger.info("[REPORTS] All report tools disabled or missing dependencies")
    
    return tools


# ============================================================
# Helper Functions
# ============================================================

def _extract_tables(markdown_content: str) -> list:
    """Extract tables from Markdown content."""
    tables = []
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        if line.startswith('|') and '|' in line:
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) > 2:  # Header + separator + data
                tables.append(table_lines)
            continue
        i += 1
    
    return tables


def _generate_chart_configs(tables: list) -> list:
    """Generate Chart.js configurations from tables.
    
    Parses each Markdown table, detects appropriate chart type,
    and generates Chart.js configuration with proper styling.
    Can return multiple charts per table if datasets need splitting.
    """
    chart_configs = []
    
    for idx, table in enumerate(tables):
        try:
            # Parse table header
            headers = [h.strip() for h in table[0].split('|')[1:-1]]
            if len(headers) < 2:
                continue
            
            # ❌ Skip tables that should NOT become charts
            # Authority Score tables should display as plain tables with red bold values
            header_text = ' '.join([h.lower() for h in headers])
            skip_keywords = [
                'authority score', 'authority', 'domain score',
                '权威分数', '域名评分'
            ]
            if any(kw in header_text for kw in skip_keywords):
                logger.debug(f"[REPORTS] Skipping chart for Authority Score table {idx}")
                continue
            
            # Parse data rows (skip separator line)
            data_rows = []
            for row in table[2:]:  # Skip separator
                cells = [c.strip() for c in row.split('|')[1:-1]]
                if cells and len(cells) >= 2:
                    data_rows.append(cells)
            
            # Need at least 2 data rows for meaningful chart
            if len(data_rows) < 2:
                continue
            
            # Detect chart type based on headers and data
            chart_type = _detect_chart_type(headers, data_rows)
            if not chart_type:
                continue
            
            # Generate Chart.js config (may return single config or list)
            config = _create_chart_config(chart_type, headers, data_rows, idx)
            if config:
                # Handle both single config and list of configs
                if isinstance(config, list):
                    chart_configs.extend(config)
                else:
                    chart_configs.append(config)
                
        except Exception as e:
            logger.debug(f"[REPORTS] Could not parse table {idx}: {e}")
    
    return chart_configs


def _detect_chart_type(headers: list, data_rows: list) -> str:
    """Detect appropriate chart type based on table structure.
    
    Detection logic (priority order):
    1. Time-series keywords in headers AND valid date patterns → Line Chart
    2. Percentage/share keywords in headers → Pie Chart (Doughnut)
    3. Metric keywords in headers → Bar Chart
    4. Has numeric data columns → Bar Chart
    5. No numeric data → None (no chart)
    """
    if len(headers) < 2 or len(data_rows) < 2:
        return None
    
    # Normalize headers for keyword matching
    header_lower = [h.lower() for h in headers]
    header_text = ' '.join(header_lower)
    
    # 1️⃣ Time-series data → Line Chart (strict validation)
    time_keywords = [
        'month', 'date', 'period', 'time', 'year', 'quarter', 'week', 'day',
        '月份', '时间', '日期', '年份', '季度',
        'mom', 'yoy', 'trend', 'history', 'historical'
    ]
    
    # First check if headers suggest time-series
    has_time_keywords = any(kw in header_text for kw in time_keywords)
    
    # Then verify first column contains actual date/time patterns
    first_col = [row[0] for row in data_rows if row]
    has_valid_dates = False
    
    if first_col and has_time_keywords:
        # Count how many rows have date-like patterns
        date_count = 0
        for val in first_col:
            val_lower = str(val).lower()
            # Check for month names
            months = ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 
                      'jul', 'aug', 'sep', 'oct', 'nov', 'dec',
                      'january', 'february', 'march', 'april', 'may', 'june',
                      'july', 'august', 'september', 'october', 'november', 'december']
            if any(month in val_lower for month in months):
                date_count += 1
                continue
            # Check for date formats (2024-01, 2024/01, Q1 2024, etc.)
            if re.match(r'\d{4}[-/]\d{2}', val_lower) or re.match(r'\d{2}[-/]\d{4}', val_lower):
                date_count += 1
                continue
            # Check for quarter patterns (Q1, Q2, etc.)
            if re.match(r'q[1-4]', val_lower):
                date_count += 1
                continue
        
        # At least 80% of rows must have valid date patterns
        if date_count >= len(first_col) * 0.8:
            has_valid_dates = True
    
    # Only use line chart if BOTH conditions are met
    if has_time_keywords and has_valid_dates:
        return 'line'
    
    # 2️⃣ Percentage/share data → Pie Chart (Doughnut)
    share_keywords = ['%', 'share', 'ratio', 'percent', 'proportion', '占比', '比例', '份额']
    if any(kw in header_text for kw in share_keywords):
        # Only use pie chart if there are few data rows (<=8)
        if len(data_rows) <= 8:
            return 'pie'
        return 'bar'  # Too many slices, use bar instead
    
    # 3️⃣ Metric keywords → Bar Chart (includes domain comparisons)
    metric_keywords = [
        'traffic', 'keywords', 'volume', 'visits', 'count', 'score', 'rank',
        'backlink', 'domain', 'revenue', 'users', 'sessions', 'pageviews',
        'organic', 'paid', 'cpc', 'ctr', 'impressions', 'clicks',
        '流量', '关键词', '排名', '得分', '用户', '访问'
    ]
    if any(kw in header_text for kw in metric_keywords):
        return 'bar'
    
    # 4️⃣ Check if there are numeric columns (fallback detection)
    numeric_col_count = 0
    for col_idx in range(1, len(headers)):
        numeric_count = 0
        for row in data_rows:
            if len(row) > col_idx:
                val = row[col_idx]
                if val and re.search(r'\d', val):
                    # Try to extract number
                    cleaned = re.sub(r'[^\d.\-]', '', val)
                    try:
                        float(cleaned)
                        numeric_count += 1
                    except ValueError:
                        pass
        # At least 50% of rows have numeric data in this column
        if numeric_count >= len(data_rows) * 0.5:
            numeric_col_count += 1
    
    if numeric_col_count > 0:
        return 'bar'
    
    # 5️⃣ No chart if no numeric data detected
    return None


def _extract_number(value: str) -> float:
    """Extract numeric value from string.
    
    Supports formats:
    - 450,000 → 450000
    - 12.5K → 12.5
    - 35% → 35
    - $8.50 → 8.50
    - +15.3% → 15.3
    - -5.2% → -5.2
    """
    if not value:
        return 0
    
    # Remove common prefixes/suffixes but keep negative sign
    cleaned = re.sub(r'[^\d.\-]', '', value)
    
    try:
        return float(cleaned) if cleaned else 0
    except ValueError:
        return 0


# Chart color palette (Minimal 2-Color: Blue #3b82f6 + Slate #475569)
CHART_COLORS = [
    {'bg': 'rgba(59, 130, 246, 0.7)', 'border': 'rgb(59, 130, 246)'},     # Primary: Blue
    {'bg': 'rgba(71, 85, 105, 0.7)', 'border': 'rgb(71, 85, 105)'},       # Secondary: Slate
    {'bg': 'rgba(59, 130, 246, 0.4)', 'border': 'rgb(59, 130, 246)'},     # Blue (lighter)
    {'bg': 'rgba(71, 85, 105, 0.4)', 'border': 'rgb(71, 85, 105)'},       # Slate (lighter)
    {'bg': 'rgba(59, 130, 246, 0.2)', 'border': 'rgb(59, 130, 246)'},     # Blue (very light)
    {'bg': 'rgba(71, 85, 105, 0.2)', 'border': 'rgb(71, 85, 105)'},       # Slate (very light)
]

PIE_COLORS = [
    'rgba(59, 130, 246, 0.85)',   # Blue
    'rgba(71, 85, 105, 0.85)',    # Slate
    'rgba(59, 130, 246, 0.6)',    # Blue (lighter)
    'rgba(71, 85, 105, 0.6)',     # Slate (lighter)
    'rgba(59, 130, 246, 0.4)',    # Blue (very light)
    'rgba(71, 85, 105, 0.4)',     # Slate (very light)
    'rgba(30, 64, 175, 0.85)',    # Darker Blue
    'rgba(51, 65, 85, 0.85)',     # Darker Slate
]


def _group_datasets_by_magnitude(datasets: list) -> list[list]:
    """Group datasets by order of magnitude to avoid scale issues.
    
    Returns list of dataset groups where each group has similar magnitudes.
    """
    if not datasets:
        return []
    
    # Calculate average magnitude for each dataset
    dataset_magnitudes = []
    for ds in datasets:
        avg_val = sum(abs(v) for v in ds['data']) / len(ds['data']) if ds['data'] else 0
        if avg_val > 0:
            magnitude = int(math.log10(avg_val))
        else:
            magnitude = 0
        dataset_magnitudes.append((ds, magnitude))
    
    # Group datasets with magnitude difference <= 2 (e.g., 100s and 10000s can go together)
    groups = []
    dataset_magnitudes.sort(key=lambda x: x[1])
    
    current_group = []
    current_mag = None
    
    for ds, mag in dataset_magnitudes:
        if current_mag is None or abs(mag - current_mag) <= 2:
            current_group.append(ds)
            if current_mag is None:
                current_mag = mag
        else:
            if current_group:
                groups.append(current_group)
            current_group = [ds]
            current_mag = mag
    
    if current_group:
        groups.append(current_group)
    
    return groups


def _generate_chart_title_from_labels(labels: list) -> str:
    """Generate a descriptive chart title from dataset labels.
    
    Analyzes dataset labels to create meaningful chart titles.
    """
    if not labels:
        return None
    
    # Common label patterns and their English titles
    label_patterns = {
        # Traffic/Keywords patterns
        ('traffic', 'kw'): 'Traffic & Keywords Trend',
        ('traffic',): 'Traffic Trend',
        ('kw', 'keyword'): 'Keyword Trend',
        ('mom', '%', 'change'): 'Month-over-Month Change',
        # Domain comparison patterns
        ('seopage', 'writesonic'): 'Domain Comparison',
        # Metric patterns
        ('organic',): 'Organic Search Metrics',
        ('paid',): 'Paid Advertising Metrics',
        ('backlink', 'ref'): 'Backlink Metrics',
        ('authority', 'score'): 'Authority Score Comparison',
        # Content patterns
        ('blog', 'product', 'landing'): 'Content Type Distribution',
        ('page', 'content'): 'Page Data',
        # Technical patterns
        ('speed', 'lcp', 'cwv'): 'Technical SEO Metrics',
    }
    
    labels_lower = ' '.join(labels).lower()
    
    for patterns, title in label_patterns.items():
        if any(p in labels_lower for p in patterns):
            return title
    
    # Fallback: extract domain names if present
    domains = []
    for label in labels:
        if '.com' in label.lower() or '.ai' in label.lower():
            domains.append(label.split()[0])
    
    if domains:
        return f"{' vs '.join(domains[:2])} Comparison"
    
    return None


def _create_chart_config(chart_type: str, headers: list, data_rows: list, idx: int) -> dict | list:
    """Create Chart.js configuration with proper styling.
    
    Returns either a single chart config dict, or a list of chart configs
    if datasets need to be split due to magnitude differences.
    """
    labels = [row[0] for row in data_rows if row]
    
    # Collect all datasets
    datasets = []
    for col_idx in range(1, len(headers)):
        data = []
        for row in data_rows:
            if len(row) > col_idx:
                data.append(_extract_number(row[col_idx]))
        
        if data and any(v != 0 for v in data):  # Skip all-zero datasets
            color = CHART_COLORS[(col_idx - 1) % len(CHART_COLORS)]
            datasets.append({
                'label': headers[col_idx],
                'data': data,
                'backgroundColor': color['bg'],
                'borderColor': color['border'],
            })
    
    if not datasets:
        return None
    
    # For pie charts, don't split
    if chart_type == 'pie':
        return {
            'id': f'chart_{idx}',
            'type': 'doughnut',
            'labels': labels,
            'datasets': datasets[:1],  # Only first dataset for pie
            'is_pie': True,
        }
    
    # Check if datasets have vastly different magnitudes
    dataset_groups = _group_datasets_by_magnitude(datasets)
    
    # If only one group or <= 3 datasets, return single chart
    if len(dataset_groups) <= 1 or len(datasets) <= 3:
        return {
            'id': f'chart_{idx}',
            'type': chart_type,
            'labels': labels,
            'datasets': datasets,
            'is_pie': False,
        }
    
    # Split into multiple charts by magnitude
    charts = []
    for group_idx, group in enumerate(dataset_groups):
        chart_id = f'chart_{idx}_{group_idx}' if group_idx > 0 else f'chart_{idx}'
        
        # Generate meaningful title from dataset labels
        group_labels = [ds['label'] for ds in group]
        chart_title = _generate_chart_title_from_labels(group_labels)
        
        # Fallback to numbered title if no pattern matched
        if not chart_title and len(dataset_groups) > 1:
            # Try to describe by magnitude scale
            avg_val = sum(sum(abs(v) for v in ds['data']) / len(ds['data']) for ds in group) / len(group)
            if avg_val >= 10000:
                chart_title = f"Large Scale Metrics (avg {avg_val:,.0f})"
            elif avg_val >= 100:
                chart_title = f"Medium Scale Metrics (avg {avg_val:,.0f})"
            else:
                chart_title = f"Small Scale Metrics (avg {avg_val:,.1f})"
        
        charts.append({
            'id': chart_id,
            'type': chart_type,
            'labels': labels,
            'datasets': group,
            'is_pie': False,
            'title': chart_title,
        })
    
    return charts


def _extract_h2_sections(content: str) -> list:
    """Extract H1 and H2 headings from HTML content for hierarchical sidebar navigation.
    
    Returns a list of sections with parent-child relationships:
    - H1 headings become parent navigation items
    - H2 headings become child navigation items under the preceding H1
    """
    import html as html_module
    sections = []
    section_idx = 0
    
    # Match both <h1> and <h2> tags
    pattern = r'<(h[12])[^>]*>([^<]+)</h[12]>'
    matches = re.finditer(pattern, content, re.IGNORECASE)
    
    current_parent = None
    
    for match in matches:
        level = match.group(1).lower()
        title = html_module.unescape(match.group(2).strip())
        section_id = f'section-{section_idx}'
        section_idx += 1
        
        if level == 'h1':
            # H1 becomes a parent navigation item
            current_parent = {
                'id': section_id,
                'title': title,
                'level': 'parent',
                'children': []
            }
            sections.append(current_parent)
        elif level == 'h2':
            # H2 becomes a child navigation item
            child = {
                'id': section_id,
                'title': title,
                'level': 'child',
                'parent': current_parent
            }
            # If there's a current parent, add as child; otherwise add as standalone
            if current_parent:
                current_parent['children'].append(child)
            else:
                # Standalone H2 (no parent H1)
                sections.append(child)
    
    return sections


def _add_section_ids_to_content(content: str, sections: list) -> str:
    """Add id attributes to H1 and H2 tags for navigation.
    
    Removes any existing IDs to prevent duplicates.
    """
    result = content
    
    # First, remove any existing id attributes from h1 and h2 tags
    result = re.sub(r'<h1([^>]*)\s+id="[^"]*"([^>]*)>', r'<h1\1\2>', result, flags=re.IGNORECASE)
    result = re.sub(r'<h2([^>]*)\s+id="[^"]*"([^>]*)>', r'<h2\1\2>', result, flags=re.IGNORECASE)
    
    # Build separate lists for H1 and H2 IDs
    h1_ids = []
    h2_ids = []
    for section in sections:
        if section.get('level') == 'parent':
            h1_ids.append(section['id'])
            for child in section.get('children', []):
                h2_ids.append(child['id'])
        else:
            h2_ids.append(section['id'])
    
    # Add IDs to H1 tags
    h1_idx = [0]  # Use list for closure mutability
    def replacer_h1(match):
        if h1_idx[0] < len(h1_ids):
            section_id = h1_ids[h1_idx[0]]
            h1_idx[0] += 1
            existing_attrs = match.group(1)
            return f'<h1{existing_attrs} id="{section_id}">'
        return match.group(0)
    
    # Add IDs to H2 tags
    h2_idx = [0]  # Use list for closure mutability
    def replacer_h2(match):
        if h2_idx[0] < len(h2_ids):
            section_id = h2_ids[h2_idx[0]]
            h2_idx[0] += 1
            existing_attrs = match.group(1)
            return f'<h2{existing_attrs} id="{section_id}">'
        return match.group(0)
    
    # Process H1 first, then H2
    result = re.sub(r'<h1([^>]*)>', replacer_h1, result, flags=re.IGNORECASE)
    result = re.sub(r'<h2([^>]*)>', replacer_h2, result, flags=re.IGNORECASE)
    
    return result


def _generate_html_template(title: str, content: str, chart_configs: list) -> str:
    """Generate complete HTML with sidebar navigation, Chart.js, and modern styling.
    
    Features:
    - Sidebar navigation based on H2 headings
    - Blue-Teal-Purple gradient theme
    - Responsive layout (desktop/tablet/mobile)
    - Interactive charts with Chart.js
    - Scroll-based navigation highlighting
    """
    # Extract sections for navigation
    sections = _extract_h2_sections(content)
    content_with_ids = _add_section_ids_to_content(content, sections)
    
    # Remove the first H1 from content (it's already in the header)
    # Find the first H1 tag and remove it along with its content
    first_h1_pattern = r'<h1[^>]*>.*?</h1>\s*'
    content_with_ids = re.sub(first_h1_pattern, '', content_with_ids, count=1, flags=re.IGNORECASE | re.DOTALL)
    
    # Generate hierarchical sidebar navigation HTML
    nav_items_html = ""
    first_section_id = None
    
    for idx, section in enumerate(sections):
        if first_section_id is None:
            first_section_id = section['id']
        
        if section.get('level') == 'parent':
            # Parent navigation item (H1)
            active_class = ' active' if section['id'] == first_section_id else ''
            nav_items_html += f'''
        <a href="#{section['id']}" class="nav-item nav-parent{active_class}" data-section="{section['id']}">
            <span class="nav-text">{section['title']}</span>
        </a>'''
            
            # Add children (H2 items)
            if section.get('children'):
                nav_items_html += '<div class="nav-children">'
                for child in section['children']:
                    nav_items_html += f'''
        <a href="#{child['id']}" class="nav-item nav-child" data-section="{child['id']}">
            <span class="nav-text">{child['title']}</span>
        </a>'''
                nav_items_html += '</div>'
        else:
            # Standalone H2 (no parent)
            active_class = ' active' if section['id'] == first_section_id else ''
            nav_items_html += f'''
        <a href="#{section['id']}" class="nav-item nav-child{active_class}" data-section="{section['id']}">
            <span class="nav-text">{section['title']}</span>
        </a>'''
    
    # Generate chart HTML and scripts
    charts_script = ""
    chart_insert_map = {}  # Map table index to chart HTML
    
    for config in chart_configs:
        if config is None:
            continue
            
        chart_id = config['id']
        chart_type = config['type']
        is_pie = config.get('is_pie', False)
        
        # Build datasets JavaScript
        if is_pie:
            # For pie charts, use first dataset with pie colors
            ds = config['datasets'][0] if config['datasets'] else None
            if ds:
                pie_colors = str(PIE_COLORS[:len(ds['data'])])
                datasets_js = f"""[{{
            label: '{ds['label']}',
            data: {ds['data']},
            backgroundColor: {pie_colors},
            borderColor: 'white',
            borderWidth: 2
        }}]"""
        else:
            # For bar/line charts
            datasets_js = "[\n"
            for ds in config['datasets']:
                tension = ", tension: 0.3, fill: true" if chart_type == 'line' else ""
                datasets_js += f"""        {{
            label: '{ds['label']}',
            data: {ds['data']},
            backgroundColor: '{ds['backgroundColor']}',
            borderColor: '{ds['borderColor']}',
            borderWidth: 2{tension}
        }},\n"""
            datasets_js += "    ]"
        
        # Chart options
        if is_pie:
            options_js = """{
        responsive: true,
        maintainAspectRatio: false,
        plugins: {
            legend: {
                position: 'right',
                labels: { color: '#374151', font: { size: 12 } }
            }
        }
    }"""
        else:
            options_js = """{
        responsive: true,
        maintainAspectRatio: false,
        plugins: {
            legend: {
                labels: { color: '#374151', font: { size: 12 } }
            }
        },
        scales: {
            x: {
                ticks: { color: '#6b7280' },
                grid: { color: 'rgba(229, 231, 235, 0.8)' }
            },
            y: {
                ticks: { color: '#6b7280' },
                grid: { color: 'rgba(229, 231, 235, 0.8)' }
            }
        }
    }"""
        
        charts_script += f"""
new Chart(document.getElementById('{chart_id}'), {{
    type: '{chart_type}',
    data: {{
        labels: {config['labels']},
        datasets: {datasets_js}
    }},
    options: {options_js}
}});
"""
        
        # Store chart HTML to insert before each table
        # Extract base table index from chart_id (e.g., "chart_0" -> 0, "chart_0_1" -> 0)
        table_idx = int(chart_id.split('_')[1])
        
        # Generate chart HTML with optional title
        chart_title_html = ""
        if config.get('title'):
            chart_title_html = f'<h4 style="margin: 0 0 16px 0; color: var(--text-secondary); font-size: 0.9rem;">{config["title"]}</h4>'
        
        # Group charts by table index
        if table_idx not in chart_insert_map:
            chart_insert_map[table_idx] = []
        
        chart_insert_map[table_idx].append(f'''
<div class="chart-section">
    {chart_title_html}
    <div class="chart-container">
        <canvas id="{chart_id}"></canvas>
    </div>
</div>''')
    
    # Insert charts before their corresponding tables
    table_pattern = r'<table>'
    table_matches = list(re.finditer(table_pattern, content_with_ids))
    
    # Insert from end to start to preserve indices
    for idx in sorted(chart_insert_map.keys(), reverse=True):
        if idx < len(table_matches):
            pos = table_matches[idx].start()
            # Join all charts for this table
            charts_html = '\n'.join(chart_insert_map[idx])
            content_with_ids = content_with_ids[:pos] + charts_html + content_with_ids[pos:]
    
    # Calculate total sections (H1 + H2)
    total_sections = sum(1 for s in sections if s.get('level') == 'parent')
    total_sections += sum(len(s.get('children', [])) for s in sections if s.get('level') == 'parent')
    total_sections += sum(1 for s in sections if s.get('level') != 'parent')
    
    # Generate the complete HTML
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
    <style>
        /* ============================================
           CSS Variables - Minimal 2-Color Theme
           Primary: Blue #3b82f6 | Secondary: Slate #475569
           ============================================ */
        :root {{
            /* Backgrounds (Gray/White) */
            --bg-primary: #fafbfc;
            --bg-secondary: #ffffff;
            --bg-sidebar: #f8f9fb;
            --bg-hover: #f3f4f6;
            
            /* Text (Black/Gray) */
            --text-primary: #1f2937;
            --text-secondary: #6b7280;
            --text-muted: #9ca3af;
            
            /* Borders (Gray) */
            --border-color: #e5e7eb;
            --border-light: #f3f4f6;
            
            /* Primary Color: Blue */
            --accent-primary: #3b82f6;
            --accent-primary-light: #eff6ff;
            
            /* Secondary Color: Slate (Gray-Blue) */
            --accent-secondary: #475569;
            --accent-secondary-light: #f1f5f9;
            
            /* Unified accent aliases for compatibility */
            --accent-blue: var(--accent-primary);
            --accent-teal: var(--accent-secondary);
            --accent-purple: var(--accent-secondary);
            --accent-blue-light: var(--accent-primary-light);
            --accent-teal-light: var(--accent-primary-light);
            --accent-purple-light: var(--accent-secondary-light);
            
            /* Status colors use primary/secondary only */
            --success: var(--accent-primary);
            --danger: var(--accent-secondary);
            --warning: var(--accent-secondary);
            
            /* Layout */
            --sidebar-width: 260px;
            
            /* Shadows */
            --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
            --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        }}
        
        /* ============================================
           Base Styles
           ============================================ */
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        html {{
            scroll-behavior: smooth;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'PingFang SC', 'Microsoft YaHei', "Segoe UI", Roboto, sans-serif;
            line-height: 1.7;
            color: var(--text-primary);
            background: var(--bg-primary);
            font-size: 14px;
            overflow-x: hidden;
        }}
        
        /* Scroll Progress Indicator */
        .scroll-indicator {{
            position: fixed;
            top: 0;
            left: 0;
            height: 3px;
            background: linear-gradient(90deg, var(--accent-primary), var(--accent-secondary));
            z-index: 200;
            transition: width 0.1s ease-out;
        }}
        
        /* ============================================
           Layout
           ============================================ */
        .layout {{
            display: flex;
            min-height: 100vh;
        }}
        
        /* ============================================
           Sidebar - Enhanced with smooth scrollbar
           ============================================ */
        .sidebar {{
            width: var(--sidebar-width);
            background: linear-gradient(180deg, var(--bg-sidebar) 0%, #fafbfc 100%);
            border-right: 1px solid var(--border-color);
            position: fixed;
            top: 0;
            left: 0;
            height: 100vh;
            overflow-y: auto;
            overflow-x: hidden;
            z-index: 100;
            box-shadow: var(--shadow-sm);
        }}
        
        /* Custom Scrollbar */
        .sidebar::-webkit-scrollbar {{
            width: 6px;
        }}
        
        .sidebar::-webkit-scrollbar-track {{
            background: transparent;
        }}
        
        .sidebar::-webkit-scrollbar-thumb {{
            background: var(--border-color);
            border-radius: 3px;
        }}
        
        .sidebar::-webkit-scrollbar-thumb:hover {{
            background: var(--text-muted);
        }}
        
        .sidebar-header {{
            padding: 24px 20px;
            border-bottom: 1px solid var(--border-color);
        }}
        
        .sidebar-brand {{
            display: flex;
            align-items: center;
            gap: 8px;
            margin-bottom: 16px;
        }}
        
        .brand-icon {{
            font-size: 20px;
            background: linear-gradient(135deg, var(--accent-primary), var(--accent-secondary));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            animation: pulse-subtle 3s ease-in-out infinite;
        }}
        
        @keyframes pulse-subtle {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.85; }}
        }}
        
        .brand-text {{
            font-size: 16px;
            font-weight: 700;
            background: linear-gradient(135deg, var(--accent-primary), var(--accent-secondary));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            letter-spacing: -0.02em;
        }}
        
        .sidebar-title {{
            font-size: 11px;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            color: var(--text-muted);
        }}
        
        .sidebar-subtitle {{
            font-size: 12px;
            color: var(--text-muted);
            margin-top: 4px;
        }}
        
        /* Navigation */
        .nav-section {{
            padding: 16px 12px;
        }}
        
        .nav-label {{
            font-size: 10px;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            color: var(--text-muted);
            padding: 0 8px;
            margin-bottom: 8px;
        }}
        
        .nav-item {{
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 10px 16px;
            margin: 2px 0;
            border-radius: 8px;
            color: var(--text-secondary);
            text-decoration: none;
            font-size: 0.875rem;
            font-weight: 500;
            transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
        }}
        
        .nav-item::before {{
            content: '';
            position: absolute;
            left: 0;
            top: 50%;
            transform: translateY(-50%);
            width: 3px;
            height: 0;
            background: linear-gradient(180deg, var(--accent-blue), var(--accent-purple));
            border-radius: 0 2px 2px 0;
            transition: height 0.2s ease;
        }}
        
        .nav-item:hover {{
            background: var(--bg-secondary);
            color: var(--text-primary);
            transform: translateX(2px);
        }}
        
        .nav-item.active {{
            background: linear-gradient(135deg, var(--accent-blue-light), var(--accent-purple-light));
            color: var(--accent-blue);
            font-weight: 600;
        }}
        
        .nav-item.active::before {{
            height: 60%;
        }}
        
        /* Parent navigation (h1 categories) */
        .nav-parent {{
            font-weight: 600;
            color: var(--text-primary);
            font-size: 0.9rem;
            margin-top: 12px;
        }}
        
        .nav-parent:first-child {{
            margin-top: 0;
        }}
        
        /* Child navigation (h2 items) */
        .nav-child {{
            padding-left: 24px;
            font-size: 0.8rem;
            color: var(--text-secondary);
        }}
        
        .nav-children {{
            margin-left: 0;
        }}
        
        .nav-child.active {{
            font-weight: 500;
        }}
        
        .nav-text {{
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
        
        /* ============================================
           Main Content
           ============================================ */
        .main-content {{
            flex: 1;
            margin-left: var(--sidebar-width);
            min-height: 100vh;
        }}
        
        .content-header {{
            position: sticky;
            top: 0;
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(8px);
            border-bottom: 1px solid var(--border-color);
            padding: 32px 48px;
            z-index: 50;
            box-shadow: var(--shadow-sm);
        }}
        
        .content-header h1 {{
            font-size: 2rem;
            font-weight: 700;
            background: linear-gradient(135deg, var(--accent-primary), var(--accent-secondary));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: 12px;
            letter-spacing: -0.03em;
            line-height: 1.2;
        }}
        
        .content-header-meta {{
            font-size: 13px;
            color: var(--text-muted);
        }}
        
        .content-body {{
            padding: 40px 48px 80px 48px;
            max-width: 1100px;
        }}
        
        /* ============================================
           Content Sections
           ============================================ */
        .content-section {{
            margin-bottom: 16px;
        }}
        
        .section-divider {{
            border: none;
            height: 1px;
            background: linear-gradient(90deg, transparent, var(--border-color), transparent);
            margin: 48px 0;
        }}
        
        .intro-section {{
            background: linear-gradient(135deg, var(--accent-primary-light) 0%, var(--accent-secondary-light) 100%);
            border-radius: 16px;
            padding: 32px;
            margin-bottom: 48px;
            border: 1px solid var(--border-light);
            box-shadow: var(--shadow-sm);
            position: relative;
            overflow: hidden;
        }}
        
        .intro-section::before {{
            content: '';
            position: absolute;
            top: -50%;
            right: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(99, 102, 241, 0.1) 0%, transparent 70%);
            animation: pulse 8s ease-in-out infinite;
            pointer-events: none;
        }}
        
        @keyframes pulse {{
            0%, 100% {{ opacity: 0.5; transform: scale(1); }}
            50% {{ opacity: 1; transform: scale(1.05); }}
        }}
        
        .intro-section > * {{
            position: relative;
            z-index: 1;
        }}
        
        /* Typography - Enhanced hierarchy */
        h1 {{
            font-size: 1.75rem;
            font-weight: 700;
            color: var(--text-primary);
            margin-top: 48px;
            margin-bottom: 24px;
            padding-bottom: 16px;
            border-bottom: 3px solid transparent;
            border-image: linear-gradient(90deg, var(--accent-primary), var(--accent-secondary)) 1;
            letter-spacing: -0.02em;
            scroll-margin-top: 100px;
        }}
        
        h1:first-child {{
            margin-top: 0;
        }}
        
        h2 {{
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--text-primary);
            margin-top: 48px;
            margin-bottom: 24px;
            padding-bottom: 16px;
            border-bottom: 3px solid transparent;
            border-image: linear-gradient(90deg, var(--accent-primary), var(--accent-secondary)) 1;
            letter-spacing: -0.02em;
            scroll-margin-top: 100px;
        }}
        
        h2:first-child {{
            margin-top: 0;
        }}
        
        h3 {{
            font-size: 1.125rem;
            font-weight: 600;
            color: var(--text-primary);
            margin-top: 32px;
            margin-bottom: 16px;
            letter-spacing: -0.01em;
        }}
        
        h4 {{
            font-size: 1rem;
            font-weight: 600;
            color: var(--text-secondary);
            margin-top: 24px;
            margin-bottom: 12px;
        }}
        
        p {{
            margin-bottom: 16px;
            color: var(--text-secondary);
            line-height: 1.8;
        }}
        
        strong {{
            color: var(--text-primary);
            font-weight: 600;
        }}
        
        /* Lists - Enhanced styling */
        ul, ol {{
            margin: 16px 0;
            padding-left: 24px;
        }}
        
        li {{
            margin-bottom: 10px;
            color: var(--text-secondary);
            line-height: 1.7;
            padding-left: 8px;
        }}
        
        ul li::marker {{
            color: var(--accent-blue);
        }}
        
        ol li::marker {{
            font-weight: 600;
            color: var(--accent-blue);
        }}
        
        /* Links - make them clearly clickable */
        a, .report-link {{
            color: var(--accent-blue);
            text-decoration: none;
            border-bottom: 1px solid transparent;
            transition: all 0.15s ease;
            cursor: pointer;
        }}
        
        a:hover, .report-link:hover {{
            color: var(--accent-purple);
            border-bottom-color: var(--accent-purple);
        }}
        
        /* Blockquotes */
        blockquote {{
            border-left: 3px solid;
            border-image: linear-gradient(180deg, var(--accent-blue), var(--accent-purple)) 1;
            padding: 16px 20px;
            margin: 16px 0;
            background: linear-gradient(135deg, var(--accent-blue-light), var(--accent-purple-light));
            border-radius: 0 8px 8px 0;
            color: var(--text-secondary);
        }}
        
        blockquote p:last-child {{
            margin-bottom: 0;
        }}
        
        /* Code - Enhanced styling */
        code {{
            background: var(--accent-purple-light);
            padding: 3px 8px;
            border-radius: 5px;
            font-family: 'SF Mono', Monaco, 'Cascadia Code', 'Courier New', monospace;
            font-size: 0.85em;
            color: var(--accent-purple);
            font-weight: 500;
        }}
        
        pre {{
            background: #1e1e2e;
            color: #cdd6f4;
            padding: 16px;
            border-radius: 8px;
            overflow-x: auto;
            margin: 16px 0;
        }}
        
        pre code {{
            background: none;
            padding: 0;
            color: inherit;
        }}
        
        /* Horizontal Rule */
        hr {{
            border: none;
            border-top: 1px solid var(--border-color);
            margin: 32px 0;
        }}
        
        /* ============================================
           Tables - Premium design
           ============================================ */
        .table-container {{
            overflow-x: auto;
            margin: 24px 0;
            border-radius: 12px;
            box-shadow: var(--shadow-sm);
            border: 1px solid var(--border-light);
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            background: var(--bg-secondary);
            font-size: 0.875rem;
        }}
        
        th {{
            background: linear-gradient(180deg, var(--bg-sidebar), var(--bg-secondary));
            color: var(--text-primary);
            font-weight: 600;
            text-align: left;
            padding: 14px 16px;
            font-size: 0.8rem;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            border-bottom: 2px solid var(--border-color);
            white-space: nowrap;
        }}
        
        td {{
            padding: 14px 16px;
            border-bottom: 1px solid var(--border-light);
            color: var(--text-secondary);
        }}
        
        /* Authority Score - Red Bold Display */
        td[data-authority-score] {{
            color: #ef4444 !important;
            font-weight: 700 !important;
            font-size: 1.1em !important;
        }}
        
        /* Highlight Authority Score column */
        th:has(+ th):nth-last-child(2):contains("Authority"),
        th:contains("Authority Score") {{
            background: linear-gradient(180deg, #fee2e2, var(--bg-secondary));
        }}
        
        tr:last-child td {{
            border-bottom: none;
        }}
        
        tr:hover td {{
            background: var(--bg-hover);
            transition: background 0.15s ease;
        }}
        
        /* Links in tables */
        .data-table a, .data-table .report-link {{
            color: var(--accent-blue);
            font-weight: 500;
        }}
        
        .data-table a:hover, .data-table .report-link:hover {{
            text-decoration: underline;
        }}
        
        /* Badges & Trends */
        .badge {{
            display: inline-block;
            padding: 4px 10px;
            border-radius: 6px;
            font-size: 0.75rem;
            font-weight: 600;
            letter-spacing: 0.3px;
        }}
        
        .badge.success {{
            background: #d1fae5;
            color: var(--success);
        }}
        
        .badge.danger {{
            background: #fee2e2;
            color: var(--danger);
        }}
        
        .badge.warning {{
            background: #fef3c7;
            color: var(--warning);
        }}
        
        /* Trend indicators with arrows */
        .trend {{
            font-weight: 600;
            font-variant-numeric: tabular-nums;
        }}
        
        .trend.up {{
            color: var(--success);
        }}
        
        .trend.up::before {{
            content: '↑ ';
            font-size: 1.1em;
        }}
        }}
        
        .trend.down {{
            color: var(--danger);
        }}
        
        .trend.down::before {{
            content: '↓ ';
            font-size: 1.1em;
        }}
        
        /* ============================================
           Charts
           ============================================ */
        /* ============================================
           Charts - Enhanced styling
           ============================================ */
        .chart-section {{
            background: var(--bg-secondary);
            border-radius: 12px;
            padding: 28px;
            margin: 32px 0;
            border: 1px solid var(--border-light);
            box-shadow: var(--shadow-md);
            transition: box-shadow 0.2s ease;
        }}
        
        .chart-section:hover {{
            box-shadow: var(--shadow-lg);
        }}
        
        .chart-container {{
            position: relative;
            height: 360px;
            width: 100%;
            margin-top: 16px;
        }}
        
        .chart-container canvas {{
            max-height: 350px !important;
        }}
        
        /* ============================================
           Footer
           ============================================ */
        .footer {{
            margin-top: 64px;
            padding-top: 32px;
            border-top: 1px solid var(--border-light);
            text-align: center;
            font-size: 0.875rem;
            color: var(--text-muted);
        }}
        
        /* ============================================
           Badges
           ============================================ */
        .badge {{
            display: inline-block;
            padding: 4px 10px;
            border-radius: 6px;
            font-size: 0.75rem;
            font-weight: 600;
            letter-spacing: 0.3px;
        }}
        
        .badge-success {{
            background: #d1fae5;
            color: var(--success);
        }}
        
        .badge-danger {{
            background: #fee2e2;
            color: var(--danger);
        }}
        
        .badge-warning {{
            background: #fef3c7;
            color: var(--warning);
        }}
        
        /* ============================================
           Responsive Design
           ============================================ */
        @media (max-width: 1024px) {{
            .sidebar {{
                width: 220px;
            }}
            
            :root {{
                --sidebar-width: 220px;
            }}
            
            .content-header,
            .content-body {{
                padding-left: 32px;
                padding-right: 32px;
            }}
        }}
        
        @media (max-width: 800px) {{
            .sidebar {{
                display: none;
            }}
            
            .main-content {{
                margin-left: 0;
            }}
            
            .content-header {{
                padding: 24px 20px;
            }}
            
            .content-header h1 {{
                font-size: 1.5rem;
            }}
            
            .content-body {{
                padding: 24px 20px 60px 20px;
            }}
            
            h2 {{
                font-size: 1.25rem;
                margin-top: 32px;
            }}
            
            .table-container {{
                margin: 16px -20px;
                border-radius: 0;
                border-left: none;
                border-right: none;
            }}
            
            .chart-container {{
                height: 280px;
                padding: 12px;
            }}
            
            table {{
                font-size: 0.8rem;
            }}
            
            th, td {{
                padding: 8px 12px;
            }}
        }}
        
        /* ============================================
           Print Styles
           ============================================ */
        @media print {{
            .sidebar {{
                display: none;
            }}
            
            .main-content {{
                margin-left: 0;
            }}
            
            .content-header {{
                position: static;
            }}
            
            .chart-container {{
                break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <div class="scroll-indicator" id="scrollIndicator"></div>
    
    <div class="layout">
        <!-- Sidebar Navigation -->
        <aside class="sidebar">
            <div class="sidebar-header">
                <div class="sidebar-brand">
                    <span class="brand-icon">✦</span>
                    <span class="brand-text">SeeNOS.ai</span>
                </div>
                <div class="sidebar-title">Table of Contents</div>
                <div class="sidebar-subtitle">{total_sections} Sections</div>
            </div>
            <nav class="nav-section">
                <div class="nav-label">Navigation</div>
                {nav_items_html}
            </nav>
        </aside>
        
        <!-- Main Content -->
        <main class="main-content">
            <header class="content-header">
                <h1>{title}</h1>
            </header>
            
            <div class="content-body">
                {content_with_ids}
                
                <div class="footer">
                    Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} | Powered by SeeNOS.ai
                </div>
            </div>
        </main>
    </div>
    
    <script>
        // ====================================
        // Scroll Progress Indicator
        // ====================================
        const scrollIndicator = document.getElementById('scrollIndicator');
        
        window.addEventListener('scroll', () => {{
            const docHeight = document.documentElement.scrollHeight - window.innerHeight;
            const scrollPercent = (window.scrollY / docHeight) * 100;
            scrollIndicator.style.width = scrollPercent + '%';
        }});
        
        // ====================================
        // Authority Score Enhancement
        // ====================================
        document.addEventListener('DOMContentLoaded', () => {{
            // Find all tables
            const tables = document.querySelectorAll('table');
            tables.forEach(table => {{
                const headers = table.querySelectorAll('th');
                let authorityColIndex = -1;
                
                // Find Authority Score column
                headers.forEach((th, index) => {{
                    const text = th.textContent.trim().toLowerCase();
                    if (text.includes('authority score') || text === 'authority') {{
                        authorityColIndex = index;
                        th.style.background = 'linear-gradient(180deg, #fee2e2, #ffffff)';
                        th.style.color = '#dc2626';
                    }}
                }});
                
                // Mark Authority Score cells
                if (authorityColIndex >= 0) {{
                    const rows = table.querySelectorAll('tbody tr');
                    rows.forEach(row => {{
                        const cells = row.querySelectorAll('td');
                        if (cells[authorityColIndex]) {{
                            const cell = cells[authorityColIndex];
                            cell.setAttribute('data-authority-score', 'true');
                            // Add visual indicator if numeric
                            const value = cell.textContent.trim();
                            if (!isNaN(value) && value !== '') {{
                                cell.innerHTML = `<strong style="color: #dc2626; font-size: 1.2em;">${{value}}</strong>`;
                            }}
                        }}
                    }});
                }}
            }});
        }});
        
        // ====================================
        // Sidebar Navigation & Scroll Highlighting
        // ====================================
        document.addEventListener('DOMContentLoaded', function() {{
            const navItems = document.querySelectorAll('.nav-item');
            const sections = document.querySelectorAll('h1[id], h2[id]');
            
            // Intersection Observer for scroll-based highlighting
            const observer = new IntersectionObserver((entries) => {{
                entries.forEach(entry => {{
                    if (entry.isIntersecting) {{
                        const id = entry.target.id;
                        navItems.forEach(item => {{
                            item.classList.toggle('active', item.getAttribute('data-section') === id);
                        }});
                    }}
                }});
            }}, {{
                threshold: 0.3,
                rootMargin: '-100px 0px -50% 0px'
            }});
            
            sections.forEach(section => observer.observe(section));
            
            // Smooth scrolling for navigation clicks
            navItems.forEach(item => {{
                item.addEventListener('click', (e) => {{
                    e.preventDefault();
                    const targetId = item.getAttribute('href').substring(1);
                    const target = document.getElementById(targetId);
                    if (target) {{
                        // Get the sticky header height
                        const header = document.querySelector('.content-header');
                        const headerHeight = header ? header.offsetHeight : 0;
                        
                        // Calculate target position with offset for sticky header
                        const targetPosition = target.getBoundingClientRect().top + window.pageYOffset - headerHeight - 20;
                        
                        // Smooth scroll to target
                        window.scrollTo({{
                            top: targetPosition,
                            behavior: 'smooth'
                        }});
                        
                        // Update active state
                        navItems.forEach(link => link.classList.remove('active'));
                        item.classList.add('active');
                    }}
                }});
            }});
        }});
        
        // ====================================
        // Chart.js Configuration
        // ====================================
        Chart.defaults.color = '#6b7280';
        Chart.defaults.borderColor = 'rgba(0, 0, 0, 0.06)';
        Chart.defaults.font.family = "'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif";
        Chart.defaults.font.size = 12;
        
        // Chart.js initialization
        {charts_script}
    </script>
</body>
</html>'''
    
    return html


def _setup_docx_styles(document):
    """Setup custom styles for Word document."""
    try:
        styles = document.styles
        
        # Code style
        if 'Code' not in [s.name for s in styles]:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(9)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
    except Exception as e:
        logger.debug(f"[REPORTS] Could not setup DOCX styles: {e}")


def _extract_table_lines(lines: list, start: int) -> tuple:
    """Extract consecutive table lines."""
    table_lines = []
    i = start
    while i < len(lines) and lines[i].startswith('|'):
        table_lines.append(lines[i])
        i += 1
    return table_lines, i


def _extract_code_block(lines: list, start: int) -> tuple:
    """Extract code block content."""
    code_lines = []
    i = start + 1
    while i < len(lines) and not lines[i].startswith('```'):
        code_lines.append(lines[i])
        i += 1
    return code_lines, i + 1


def _add_table_to_docx(document, table_lines: list):
    """Add a table to Word document."""
    try:
        # Parse table
        headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
        data_rows = []
        for row in table_lines[2:]:  # Skip separator
            cells = [c.strip() for c in row.split('|')[1:-1]]
            if cells:
                data_rows.append(cells)
        
        if not data_rows:
            return
        
        # Create table
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # Add data rows
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_data
        
        document.add_paragraph()  # Add spacing
    except Exception as e:
        logger.debug(f"[REPORTS] Error adding table to DOCX: {e}")


def _clean_markdown(text: str) -> str:
    """Remove Markdown formatting from text."""
    # Remove bold/italic
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Remove links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

